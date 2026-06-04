from pathlib import Path
from datetime import datetime
from docx import Document


class FileSaver:
    def __init__(self, base_dir: str = "data/output"):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def _sanitize_name(self, name: str) -> str:
        invalid_chars = '<>:"/\\|?*'
        sanitized = name
        for ch in invalid_chars:
            sanitized = sanitized.replace(ch, "_")
        sanitized = sanitized.strip()
        return sanitized[:50] if len(sanitized) > 50 else sanitized

    def create_run_folder(self, topic: str = "") -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = self._sanitize_name(topic) if topic else "untitled"
        run_dir = self.base_dir / f"{timestamp}_{safe_topic}"
        run_dir.mkdir(parents=True, exist_ok=True)
        return run_dir

    def save_docx(self, folder: Path, filename: str, content: str, title: str = ""):
        file_path = folder / filename
        document = Document()

        if title:
            document.add_heading(title, level=1)

        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if line:
                document.add_paragraph(line)

        document.save(file_path)