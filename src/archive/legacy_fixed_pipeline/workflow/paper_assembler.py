from pathlib import Path
from datetime import datetime
from docx import Document

from src.input_loader.file_reader import FileReader
from src.utils.config_loader import ConfigLoader
from src.utils.logger import get_logger
from src.utils.section_text_cleaner import SectionTextCleaner
from src.formatting.ruledoc_adapter import RuleDocAdapter


class PaperAssembler:
    def __init__(self):
        self.config_loader = ConfigLoader()
        self.logger = get_logger("PaperAssembler")
        self.file_reader = FileReader()
        self.text_cleaner = SectionTextCleaner()

        paths_config = self.config_loader.load_paths_config()
        runtime_config = self.config_loader.load_runtime_config()

        self.sections_dir = Path(paths_config["output"]["sections_dir"])
        self.assembled_dir = Path(paths_config["output"]["assembled_dir"])
        self.assembled_dir.mkdir(parents=True, exist_ok=True)

        self.enable_ruledoc = runtime_config["ruledoc"]["enabled"]
        ruledoc_rule_name = runtime_config["ruledoc"]["rule_name"]
        self.ruledoc_adapter = RuleDocAdapter(rule_name=ruledoc_rule_name)

    def _sanitize_name(self, name: str) -> str:
        invalid_chars = '<>:"/\\|?*'
        sanitized = name
        for ch in invalid_chars:
            sanitized = sanitized.replace(ch, "_")
        sanitized = sanitized.strip()
        return sanitized[:50] if len(sanitized) > 50 else sanitized

    def _create_run_folder(self, topic: str) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = self._sanitize_name(topic) if topic else "untitled"
        run_dir = self.assembled_dir / f"{timestamp}_{safe_topic}"
        run_dir.mkdir(parents=True, exist_ok=True)
        return run_dir

    def _find_latest_file(self, base_dir: Path, filename: str, topic: str = "") -> Path:
        if not base_dir.exists():
            raise FileNotFoundError(f"目录不存在: {base_dir}")

        matched_files = list(base_dir.rglob(filename))

        if topic:
            matched_files = [p for p in matched_files if topic in str(p.parent)]

        if not matched_files:
            raise FileNotFoundError(f"未找到文件: {filename}，搜索目录: {base_dir}")

        matched_files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
        return matched_files[0]

    def _add_section(self, document: Document, heading: str, body_text: str):
        document.add_heading(heading, level=1)

        lines = body_text.split("\n")
        for line in lines:
            line = line.strip()
            if line:
                document.add_paragraph(line)

    def assemble_from_latest_sections(self, topic: str) -> dict:
        self.logger.info(f"开始组装全文 | topic={topic}")

        introduction_path = self._find_latest_file(
            self.sections_dir, "04_introduction_polished.docx", topic=topic
        )
        related_work_path = self._find_latest_file(
            self.sections_dir, "06_related_work_polished.docx", topic=topic
        )
        methodology_path = self._find_latest_file(
            self.sections_dir, "08_methodology_polished.docx", topic=topic
        )
        experiment_path = self._find_latest_file(
            self.sections_dir, "10_experiment_polished.docx", topic=topic
        )
        conclusion_path = self._find_latest_file(
            self.sections_dir, "12_conclusion_polished.docx", topic=topic
        )

        introduction_text = self.file_reader.read_input(str(introduction_path))
        related_work_text = self.file_reader.read_input(str(related_work_path))
        methodology_text = self.file_reader.read_input(str(methodology_path))
        experiment_text = self.file_reader.read_input(str(experiment_path))
        conclusion_text = self.file_reader.read_input(str(conclusion_path))

        return self.assemble_from_texts(
            topic=topic,
            introduction_text=introduction_text,
            related_work_text=related_work_text,
            methodology_text=methodology_text,
            experiment_text=experiment_text,
            conclusion_text=conclusion_text,
            source_paths={
                "introduction": str(introduction_path),
                "related_work": str(related_work_path),
                "methodology": str(methodology_path),
                "experiment": str(experiment_path),
                "conclusion": str(conclusion_path),
            }
        )

    def assemble_from_texts(
        self,
        topic: str,
        introduction_text: str,
        related_work_text: str,
        methodology_text: str,
        experiment_text: str,
        conclusion_text: str,
        source_paths: dict | None = None
    ) -> dict:
        run_dir = self._create_run_folder(topic)

        final_docx_path = run_dir / "final_paper.docx"
        final_ruledoc_path = run_dir / "final_paper_ruledoc.docx"

        # 在组装前先清洗章节文本，避免 Markdown 标题、过程标签等污染最终成稿
        introduction_text = self.text_cleaner.clean(introduction_text)
        related_work_text = self.text_cleaner.clean(related_work_text)
        methodology_text = self.text_cleaner.clean(methodology_text)
        experiment_text = self.text_cleaner.clean(experiment_text)
        conclusion_text = self.text_cleaner.clean(conclusion_text)

        document = Document()
        document.add_heading(topic, level=0)

        self._add_section(document, "1 引言", introduction_text)
        self._add_section(document, "2 相关工作", related_work_text)
        self._add_section(document, "3 方法设计", methodology_text)
        self._add_section(document, "4 实验与结果分析", experiment_text)
        self._add_section(document, "5 结论", conclusion_text)

        document.save(final_docx_path)
        self.logger.info(f"全文组装完成 | output={final_docx_path}")

        ruledoc_output = ""
        if self.enable_ruledoc:
            self.ruledoc_adapter.format_docx(
                input_path=str(final_docx_path),
                output_path=str(final_ruledoc_path)
            )
            ruledoc_output = str(final_ruledoc_path)
            self.logger.info(f"全文 RuleDoc 格式修正完成 | output={final_ruledoc_path}")

        return {
            "topic": topic,
            "run_dir": str(run_dir),
            "final_docx_path": str(final_docx_path),
            "final_ruledoc_path": ruledoc_output,
            "source_paths": source_paths or {}
        }