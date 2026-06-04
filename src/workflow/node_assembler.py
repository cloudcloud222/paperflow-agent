import json
from pathlib import Path
from datetime import datetime
from docx import Document

from src.utils.config_loader import ConfigLoader
from src.utils.logger import get_logger
from src.utils.section_text_cleaner import SectionTextCleaner
from src.formatting.ruledoc_adapter import RuleDocAdapter


class NodeAssembler:
    def __init__(self):
        self.config_loader = ConfigLoader()
        self.logger = get_logger("NodeAssembler")
        self.text_cleaner = SectionTextCleaner()

        paths_config = self.config_loader.load_paths_config()
        runtime_config = self.config_loader.load_runtime_config()

        self.schema_dir = Path(paths_config["output"]["outline_schema_dir"])
        self.node_assembled_dir = Path(paths_config["output"]["node_assembled_dir"])
        self.node_assembled_dir.mkdir(parents=True, exist_ok=True)

        self.enable_ruledoc = runtime_config["ruledoc"]["enabled"]
        ruledoc_rule_name = runtime_config["ruledoc"]["rule_name"]
        self.ruledoc_adapter = RuleDocAdapter(rule_name=ruledoc_rule_name)

    def _sanitize_name(self, name: str) -> str:
        invalid_chars = '<>:"/\\|?*'
        sanitized = name
        for ch in invalid_chars:
            sanitized = sanitized.replace(ch, "_")
        sanitized = sanitized.strip()
        return sanitized[:50] if len(sanitized) > 50 else sanitized

    def _create_run_folder(self, topic: str) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = self._sanitize_name(topic) if topic else "untitled"
        run_dir = self.node_assembled_dir / f"{timestamp}_{safe_topic}"
        run_dir.mkdir(parents=True, exist_ok=True)
        return run_dir

    def _find_latest_schema_file(self, topic: str) -> Path:
        if not self.schema_dir.exists():
            raise FileNotFoundError(f"结构化总纲目录不存在: {self.schema_dir}")

        matched_files = list(self.schema_dir.glob("*.json"))
        matched_files = [p for p in matched_files if topic in p.name]

        if not matched_files:
            raise FileNotFoundError(f"未找到当前主题对应的结构化总纲 JSON: topic={topic}")

        matched_files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
        return matched_files[0]

    def _node_heading_text(self, node: dict) -> str:
        number = node.get("number", "").strip()
        title = node.get("title", "").strip()
        return f"{number} {title}".strip()

    def _add_node_recursive(
        self,
        document: Document,
        node: dict,
        node_text_map: dict,
        include_empty_placeholder: bool = False
    ):
        heading_text = self._node_heading_text(node)
        level = int(node.get("level", 1))
        heading_level = min(max(level, 1), 9)

        document.add_heading(heading_text, level=heading_level)

        node_id = node.get("id", "")
        body_text = node_text_map.get(node_id, "").strip()

        if body_text:
            cleaned_text = self.text_cleaner.clean(body_text)
            for line in cleaned_text.split("\n"):
                line = line.strip()
                if line:
                    document.add_paragraph(line)
        elif include_empty_placeholder:
            document.add_paragraph("[该节点正文暂未生成]")

        for child in node.get("children", []):
            self._add_node_recursive(
                document=document,
                node=child,
                node_text_map=node_text_map,
                include_empty_placeholder=include_empty_placeholder
            )

    def assemble_from_schema_dict(
        self,
        schema: dict,
        node_text_map: dict,
        include_empty_placeholder: bool = False
    ) -> dict:
        topic = schema.get("paper_title", "untitled")
        run_dir = self._create_run_folder(topic)

        final_docx_path = run_dir / "node_assembled_paper.docx"
        final_ruledoc_path = run_dir / "node_assembled_paper_ruledoc.docx"
        node_map_path = run_dir / "node_text_map_used.json"

        node_map_path.write_text(
            json.dumps(node_text_map, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

        document = Document()
        document.add_heading(topic, level=0)

        writing_goal = schema.get("writing_goal", "").strip()
        if writing_goal:
            document.add_paragraph(f"写作总目标：{writing_goal}")

        for chapter in schema.get("chapters", []):
            self._add_node_recursive(
                document=document,
                node=chapter,
                node_text_map=node_text_map,
                include_empty_placeholder=include_empty_placeholder
            )

        document.save(final_docx_path)
        self.logger.info(f"节点递归组装完成 | output={final_docx_path}")

        ruledoc_output = ""
        if self.enable_ruledoc:
            self.ruledoc_adapter.format_docx(
                input_path=str(final_docx_path),
                output_path=str(final_ruledoc_path)
            )
            ruledoc_output = str(final_ruledoc_path)
            self.logger.info(f"节点递归组装 RuleDoc 完成 | output={final_ruledoc_path}")

        return {
            "topic": topic,
            "run_dir": str(run_dir),
            "final_docx_path": str(final_docx_path),
            "final_ruledoc_path": ruledoc_output,
            "node_text_map_path": str(node_map_path)
        }

    def assemble_from_schema_file(
        self,
        schema_json_path: str,
        node_text_map: dict,
        include_empty_placeholder: bool = False
    ) -> dict:
        path = Path(schema_json_path)
        if not path.exists():
            raise FileNotFoundError(f"结构化总纲 JSON 不存在: {path}")

        schema = json.loads(path.read_text(encoding="utf-8"))
        return self.assemble_from_schema_dict(
            schema=schema,
            node_text_map=node_text_map,
            include_empty_placeholder=include_empty_placeholder
        )

    def assemble_latest_schema(
        self,
        topic: str,
        node_text_map: dict,
        include_empty_placeholder: bool = False
    ) -> dict:
        schema_path = self._find_latest_schema_file(topic)
        self.logger.info(f"开始按最新结构化总纲递归组装 | topic={topic} | schema={schema_path}")
        return self.assemble_from_schema_file(
            schema_json_path=str(schema_path),
            node_text_map=node_text_map,
            include_empty_placeholder=include_empty_placeholder
        )