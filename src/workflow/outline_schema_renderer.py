import json
from pathlib import Path
from datetime import datetime
from docx import Document

from src.utils.config_loader import ConfigLoader
from src.utils.logger import get_logger


class OutlineSchemaRenderer:
    def __init__(self):
        self.config_loader = ConfigLoader()
        self.logger = get_logger("OutlineSchemaRenderer")

        paths_config = self.config_loader.load_paths_config()
        self.schema_dir = Path(paths_config["output"]["outline_schema_dir"])
        self.rendered_dir = Path(paths_config["output"]["outline_schema_rendered_dir"])
        self.rendered_dir.mkdir(parents=True, exist_ok=True)

    def _sanitize_name(self, name: str) -> str:
        invalid_chars = '<>:"/\\|?*'
        sanitized = name
        for ch in invalid_chars:
            sanitized = sanitized.replace(ch, "_")
        sanitized = sanitized.strip()
        return sanitized[:50] if len(sanitized) > 50 else sanitized

    def _create_run_folder(self, topic: str) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = self._sanitize_name(topic) if topic else "untitled"
        run_dir = self.rendered_dir / f"{timestamp}_{safe_topic}"
        run_dir.mkdir(parents=True, exist_ok=True)
        return run_dir

    def _find_latest_schema_file(self, topic: str) -> Path:
        if not self.schema_dir.exists():
            raise FileNotFoundError(f"结构化总纲目录不存在: {self.schema_dir}")

        matched_files = list(self.schema_dir.glob("*.json"))
        matched_files = [p for p in matched_files if topic in p.name]

        if not matched_files:
            raise FileNotFoundError(f"未找到当前主题对应的结构化总纲 JSON: topic={topic}")

        matched_files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
        return matched_files[0]

    def _render_node_to_lines(self, node: dict, lines: list[str]):
        number = node.get("number", "")
        title = node.get("title", "")
        goal = node.get("goal", "")
        key_points = node.get("key_points", [])
        node_type = node.get("node_type", "")
        children = node.get("children", [])

        title_line = f"{number} {title}".strip()
        lines.append(title_line)

        if node_type:
            lines.append(f"  节点类型：{node_type}")

        if goal:
            lines.append(f"  本节点目标：{goal}")

        if key_points:
            lines.append("  核心要点：")
            for point in key_points:
                lines.append(f"    - {point}")

        lines.append("")

        for child in children:
            self._render_node_to_lines(child, lines)

    def _render_node_to_docx(self, document: Document, node: dict):
        number = node.get("number", "")
        title = node.get("title", "")
        level = int(node.get("level", 1))
        goal = node.get("goal", "")
        key_points = node.get("key_points", [])
        node_type = node.get("node_type", "")
        children = node.get("children", [])

        heading_text = f"{number} {title}".strip()
        heading_level = min(max(level, 1), 9)
        document.add_heading(heading_text, level=heading_level)

        if node_type:
            document.add_paragraph(f"节点类型：{node_type}")

        if goal:
            document.add_paragraph(f"本节点目标：{goal}")

        if key_points:
            document.add_paragraph("核心要点：")
            for point in key_points:
                document.add_paragraph(point, style="List Bullet")

        for child in children:
            self._render_node_to_docx(document, child)

    def render_from_schema(self, schema: dict) -> dict:
        topic = schema.get("paper_title", "untitled")
        run_dir = self._create_run_folder(topic)

        txt_path = run_dir / "outline_schema_rendered.txt"
        docx_path = run_dir / "outline_schema_rendered.docx"

        # 渲染 txt
        lines = []
        lines.append(f"论文题目建议：{schema.get('paper_title', '')}")
        lines.append(f"写作总目标：{schema.get('writing_goal', '')}")
        lines.append("")

        for chapter in schema.get("chapters", []):
            self._render_node_to_lines(chapter, lines)

        txt_content = "\n".join(lines).strip()
        txt_path.write_text(txt_content, encoding="utf-8")

        # 渲染 docx
        document = Document()
        document.add_heading(schema.get("paper_title", "论文结构化总纲"), level=0)

        writing_goal = schema.get("writing_goal", "")
        if writing_goal:
            document.add_paragraph(f"写作总目标：{writing_goal}")

        for chapter in schema.get("chapters", []):
            self._render_node_to_docx(document, chapter)

        document.save(docx_path)

        self.logger.info(f"结构化总纲渲染完成 | topic={topic} | txt={txt_path} | docx={docx_path}")

        return {
            "topic": topic,
            "run_dir": str(run_dir),
            "txt_path": str(txt_path),
            "docx_path": str(docx_path)
        }

    def render_from_json_file(self, json_path: str) -> dict:
        path = Path(json_path)
        if not path.exists():
            raise FileNotFoundError(f"JSON 文件不存在: {path}")

        schema = json.loads(path.read_text(encoding="utf-8"))
        return self.render_from_schema(schema)

    def render_latest_schema(self, topic: str) -> dict:
        latest_schema_path = self._find_latest_schema_file(topic)
        self.logger.info(f"开始渲染最新结构化总纲 | topic={topic} | source={latest_schema_path}")
        return self.render_from_json_file(str(latest_schema_path))